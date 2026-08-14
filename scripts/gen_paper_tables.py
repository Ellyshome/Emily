"""生成论文表格 docx —— 从图表制作文件提取 9 张表，排版为 Word 表格。

输出：docs/论文类/附图类/论文表格.docx
样式：宋体标题 + 表头加粗底纹，符合 GB 规范（表题在上方）。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 表格数据（直接取自图表制作文件）──
# 每项：(编号, 标题, 表头列表, 数据行列表[list of list])
TABLES = [
    (
        "表 1-1",
        "四重困难 × 设计原则对应表",
        ["组织视角的结构性困难", "对应设计原则", "解答逻辑"],
        [
            ["行为改变成本高", "P3 自然语言信息接入", "低学习门槛的 IM 入口消解习惯改变阻力"],
            ["数据冷启动死循环", "P5 陪跑式信息采集", "不要求刻意贡献，系统在运行中自然沉淀"],
            ["组织风险顾虑", "P1 弱人工智能工具化定位", "只陪跑不卡流程，签名认证式人机边界"],
            ["价值归属模糊", "P2 团队级训练主体", "知识沉淀到团队而非个人，价值对组织可见"],
            ["（技术使能层）", "P4 渐进式知识发现", "上下文窗口约束下的工程方案"],
        ],
    ),
    (
        "表 2-1",
        "研究分类框架表",
        ["研究类型", "AI 服务对象", "数据来源", "使用方式", "主要局限"],
        [
            ["单点专业工具", "个人", "单次输入", "即时调用", "缺乏长期项目记忆"],
            ["RAG 问答", "个人 / 团队", "文档库", "被动查询", "缺乏过程数据"],
            ["BIM / 数字孪生", "项目", "结构化工程数据", "平台操作", "信息录入成本较高"],
            ["Agent 系统", "任务", "工具 / 数据库", "任务执行", "多为任务中心"],
            ["本文系统", "团队", "日常过程信息", "长期陪跑", "尚缺大规模组织验证"],
        ],
    ),
    (
        "表 2-2",
        "本文与既有路径对比表",
        ["维度", "传统数字化工具（MIS/BIM）", "通用 AI 助手", "行业 AI 单点工具 / 平台问答", "本文系统"],
        [
            ["信息入口", "表单 / 模型（需培训）", "网页/APP 对话", "平台界面 / 专业软件", "IM 自然语言（低门槛）"],
            ["基本单位", "流程", "个人", "任务", "团队"],
            ["知识主体", "企业（格式固化）", "个人", "平台 / 个人", "团队（结构化组织资产）"],
            ["信息汇聚", "汇总后处理", "无", "查询式", "全量陪跑采集"],
            ["人机边界", "流程控制", "无约束", "系统内约束", "签名认证式"],
            ["能力生长", "固定", "固定", "知识库更新", "受控持续演进"],
            ["事实校验", "无（单一来源信息）", "无", "无", "相对独立的信息记录与交叉溯源"],
        ],
    ),
    (
        "表 3-1",
        "问题—约束—原则映射模型",
        ["结构性问题", "设计约束", "对应设计原则"],
        [
            ["信息录入困难", "操作门槛", "P3 自然语言信息接入"],
            ["数据冷启动", "使用前没有数据", "P5 陪跑式信息采集"],
            ["AI 越权风险", "决策责任不清", "P1 弱人工智能工具化定位"],
            ["知识随人流失", "知识主体个人化", "P2 团队级训练主体"],
            ["知识规模扩大", "上下文 / 注意力限制", "P4 渐进式知识发现"],
        ],
    ),
    (
        "表 3-2",
        "五原则结构关系表",
        ["结构层", "原则", "解决的问题"],
        [
            ["输入侧", "P3 自然语言信息接入", "信息怎么进来"],
            ["输入侧", "P5 陪跑式信息采集", "信息如何持续积累"],
            ["处理侧", "P4 渐进式知识发现", "知识与能力如何被组织与调用"],
            ["边界侧", "P1 弱人工智能工具化定位", "AI 能做什么、不能做什么"],
            ["边界侧", "P2 团队级训练主体", "知识归谁所有、为谁沉淀"],
        ],
    ),
    (
        "表 3-3",
        "术语对照表",
        ["术语", "学术化定义"],
        [
            ["工作任务单元（WorkItem）", "一条用户消息拆解出的独立业务意图执行单元"],
            ["三书", "团队型智能体的上下文组织机制：环境上下文（项目态势）+ 组织上下文（规则）+ 能力上下文（系统能力）"],
            ["签名认证制", "信息自由录入、权限人签证认领后方视为有效的记录生效机制"],
            ["陪跑式信息采集", "系统在日常运行中自然采集过程信息、不依赖用户刻意填报的数据积累方式"],
            ["渐进式知识发现", "“目录索引先行、按需加载全文”的知识与能力加载策略"],
            ["受控持续演进", "人工确认约束下的系统持续观察、沉淀经验与能力增量迭代机制"],
            ["可追溯记录链", "由不同角色、不同时间、不同渠道留痕串成的可对照任务推进记录"],
        ],
    ),
    (
        "表 4-1",
        "设计决策与原则映射表",
        ["架构层次", "关键设计决策", "对应原则", "决策核心理由"],
        [
            ["接入层", "双容器结构，薄插件+独立内核", "P3", "IM 为入口但不为 IM 所绑定"],
            ["会话层", "会话智能体：三书认知基石（态势书+能力书；审核智能体携规则书）", "P4", "组织规则+项目态势+能力边界，按需注入"],
            ["会话层", "会话智能体：目录索引注入，只组织不执行", "P4", "上下文窗口约束与注意力保持"],
            ["会话层", "资源范围与成果约束随任务下放", "P1", "任务执行边界由会话层预先界定"],
            ["会话层", "权限快照一次性注入、只读", "P1", "杜绝智能体越权篡改"],
            ["执行层", "任务智能体：单一任务聚焦的小型驾驭层", "P4", "不掌握无关项目信息，避免知识污染"],
            ["执行层", "结构化输出优先模式", "P1", "工具选择应固定于流程而非 LLM 临场推理"],
            ["执行层", "审核智能体：规则书+鉴权+成果约束三重灌注", "P1", "AI 产出经独立核验关卡后方达用户"],
            ["执行层", "错误分析纠错闭环", "P1", "AI 不确定性环节的自纠能力"],
            ["知识层", "SOP 集合持续自更新", "P2", "经验随团队沉淀、优胜劣汰"],
            ["知识层", "专家智能体：集团管控标准进化体", "P4", "管控手册提炼为按需加载的领域专家"],
            ["数据层", "签名认证制事件", "P1", "只陪跑不卡流程的人机边界"],
            ["数据层", "声明式拦截器", "P1", "横切逻辑与业务解耦"],
            ["数据层", "全景节点+项目日志", "P5", "陪跑式采集的数据底座"],
            ["演进层", "人主导的经验结构化（增量定义 SOP/节点模板）", "P2", "经验随团队沉淀、跨项目复用"],
            ["演进层", "系统自我观察与迭代（洞察→规则→补丁闭环）", "P1", "人批准、系统执行、全程留痕"],
        ],
    ),
    (
        "表 5-1",
        "功能验证结果汇总表",
        ["验证维度", "场景数", "通过", "通过率"],
        [
            ["SOP 意图识别与业务路由", "15", "15", "100%"],
            ["结构化参数提取与工具调用", "12", "11", "91.7%"],
            ["数据查询与结果过滤", "8", "8", "100%"],
            ["权限分级控制", "6", "5", "83.3%"],
            ["文件管理", "11", "9", "81.8%"],
            ["多轮对话与上下文保持", "4", "4", "100%"],
            ["复合请求拆解与确认闭环", "6", "6", "100%"],
            ["合计", "62", "58", "93.5%"],
        ],
    ),
    (
        "表 5-2",
        "设计原则验证矩阵",
        ["设计原则", "验证场景", "验证结论"],
        [
            ["P1 弱人工智能工具化定位", "场景三", "专家智能体输出量化评分报告，最终决策权归人"],
            ["P2 团队级训练主体", "场景一", "越级查询跨管理层级直接透传，无需中间汇报"],
            ["P3 自然语言信息接入", "场景一、二、三", "全程自然语言交互，无表单、无预配置报表"],
            ["P4 渐进式知识发现", "场景三", "通用智能体按需加载专家手册切换角色，非全量注入"],
            ["P5 陪跑式信息采集", "场景二", "日报由系统日常采集的事件/任务动态聚合，非预配置模板"],
        ],
    ),
    (
        "表 5-3",
        "可追溯记录链对比表（项目 A vs 项目 B）",
        ["维度", "项目 B（传统模式，未使用 Emily）", "项目 A（使用 Emily）"],
        [
            ["指令下达", "集团微信群（滚动淹没）", "任务下发 TSK + 全景节点，均留痕"],
            ["执行留痕", "工长微信口头汇报 + 拍照", "完工上报事件（执行人，附照片地址）"],
            ["记录来源", "记录人=执行人=上报人，同一只手", "执行人 ≠ 认证人 ≠ 验收人，多角色"],
            ["时间可对照", "否（各说各话，无汇聚点）", "是（7/15 下发 → 7/18 上报 → 7/19 验收）"],
            ["交叉验证", "无从谈起（孤岛）", "可追溯记录链完整，可组合交叉验证"],
            ["领导核查方式", "逐人电话 / 翻群 / 找台账", "一次自然语言查询"],
        ],
    ),
    (
        "表 5-4",
        "连续过程案例五阶段表",
        ["阶段", "时间", "角色", "动作", "事件类型", "关联关系"],
        [
            ["1 问题发现", "Day 1", "项目经理", "记录铺装色差问题", "quality_issue", "锚点（无前向关联）"],
            ["2 原因分析", "Day 3", "设计负责人", "确认材料批次差异", "decision", "关联 E1"],
            ["3 处理方案", "Day 5", "工程负责人", "供应商重新送样", "quality_issue", "关联 E1"],
            ["4 结果确认", "Day 7", "项目经理", "确认定样、闭环", "decision", "关联 E1 + E3（自动补全）"],
            ["5 回溯查询", "Day 8+", "项目总监（未参与）", "自然语言回溯全程", "—", "查询"],
        ],
    ),
    (
        "表 5-5",
        "事件关联链表",
        ["事件编号", "标题", "事件类型", "关联事件（related_event_ids）", "记录人", "确认人"],
        [
            ["E1", "铺装颜色与效果图差异大", "quality_issue", "[]", "张正宏", "张正宏"],
            ["E2", "铺装色差系材料批次", "decision", "[E1]", "赵明远", "赵明远"],
            ["E3", "铺装样品重新送样", "quality_issue", "[E1]", "李景利", "李景利"],
            ["E4", "样板区铺装定样", "decision", "[E1, E3]", "张正宏", "张正宏"],
        ],
    ),
]


def set_cell_font(cell, text, bold=False, size=9, align=None, header=False):
    """设置单元格文字与格式。"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if header:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # 垂直居中
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def set_cell_shading(cell, color="D9E2F3"):
    """设置单元格底纹。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_table(doc, num, title, headers, rows):
    """添加一张带表题的表格。"""
    # 表题（表题在上方，GB 规范）
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"{num}  {title}")
    title_run.font.name = "宋体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True

    # 表格
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_font(cell, h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
        set_cell_shading(cell, "D9E2F3")

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            # 数字/短字段居中，长文本左对齐
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(cell, str(val), bold=False, size=9, align=align)

    # 表格后空一行
    doc.add_paragraph()


def main():
    doc = Document()

    # 页面设置：A4，正常页边距
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 文档标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("论文表格汇总")
    tr.font.name = "宋体"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    tr.font.size = Pt(16)
    tr.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("《面向地产项目团队的信息协同型 AI 智能体系统设计与实践》")
    sr.font.name = "宋体"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    sr.font.size = Pt(10.5)

    doc.add_paragraph()

    for num, title_text, headers, rows in TABLES:
        add_table(doc, num, title_text, headers, rows)

    out = "docs/论文类/附图类/论文表格.docx"
    # 目标文件被占用时（如已在 Word 打开），回退到带时间戳的新文件名
    try:
        doc.save(out)
    except PermissionError:
        import time
        out = f"docs/论文类/附图类/论文表格_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(out)
    print(f"已生成 {out}，共 {len(TABLES)} 张表")


if __name__ == "__main__":
    main()
