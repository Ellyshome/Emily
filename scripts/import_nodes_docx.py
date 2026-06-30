#!/usr/bin/env python3
"""全景节点图 V2 DOCX 文档解析导入工具 —— 需求文档 §5.1（Phase 2）。

管线：DOCX 文件 → python-docx 提取文本+表格 → LLM chat_json 结构化提取 → 节点数据 → NodeService 导入

用法：
    uv run python scripts/import_nodes_docx.py <docx文件路径> --project-id <项目ID> [--dry-run]

依赖：
    pip install python-docx  # DOCX 文本提取
"""

from __future__ import annotations

import argparse
import asyncio
import json
import sys
from pathlib import Path

# 复用 PDF 解析器的 LLM 提取 schema 和 prompt（如果可用）
try:
    from scripts.import_nodes_pdf import (
        NODE_EXTRACTION_SCHEMA,
        EXTRACTION_PROMPT,
        extract_nodes_with_llm,
    )
except ImportError:
    NODE_EXTRACTION_SCHEMA = {
        "type": "object",
        "properties": {
            "nodes": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "node_id": {"type": "string"},
                        "node_name": {"type": "string"},
                        "stage_id": {"type": "integer"},
                        "deadline": {"type": "string"},
                        "owner_dept_id": {"type": "string"},
                        "parent_node_id": {"type": "string"},
                        "deliverables": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "deliverable_name": {"type": "string"},
                                    "target_amount": {"type": "number"},
                                    "unit": {"type": "string"},
                                    "is_required": {"type": "boolean"},
                                },
                                "required": ["deliverable_name", "target_amount", "unit"],
                            },
                        },
                        "dependencies": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "depends_on_deliverable_id": {"type": "string"},
                                    "weight": {"type": "number"},
                                },
                                "required": ["depends_on_deliverable_id"],
                            },
                        },
                    },
                    "required": ["node_id", "node_name"],
                },
            }
        },
        "required": ["nodes"],
    }

    EXTRACTION_PROMPT = """你是一个工程项目计划解析专家。请从以下文档文本中提取所有项目节点信息。

提取规则：
1. 识别每个工作项/节点的：编号、名称、截止时间、主责部门
2. 识别成果物（交付文件/里程碑产出）：名称、数量、单位
3. 识别前置依赖关系：哪个节点需要哪些前置文件
4. 如果文档中字段缺失，不要编造——留空或使用默认值

文档文本：
---
{text}
---

请以 JSON 格式返回提取的节点列表。"""

    async def extract_nodes_with_llm(text: str) -> list[dict]:
        """本地 stub，生产环境走 PDF 脚本的同名函数。"""
        return []


def parse_args():
    parser = argparse.ArgumentParser(description="全景节点图 DOCX 批量导入")
    parser.add_argument("path", type=str, help="DOCX 文件路径或目录路径")
    parser.add_argument("--project-id", type=str, required=True, help="目标项目ID")
    parser.add_argument("--creator-id", type=str, default="docx-import", help="创建人ID")
    parser.add_argument("--dry-run", action="store_true", help="仅解析不导入")
    parser.add_argument("--extract-only", action="store_true", help="仅提取文本不调用 LLM")
    return parser.parse_args()


def extract_text_from_docx(filepath: str) -> str:
    """使用 python-docx 提取 DOCX 纯文本 + 表格。"""
    try:
        from docx import Document
    except ImportError:
        print("错误：需要 python-docx 库。请执行: pip install python-docx")
        sys.exit(1)

    doc = Document(filepath)
    parts = []

    # 提取段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            if "Heading" in style or "标题" in style:
                parts.append(f"# {para.text.strip()}")
            else:
                parts.append(para.text.strip())

    # 提取表格
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append(f"\n[表格 {i+1}]\n" + "\n".join(rows))

    return "\n".join(parts)


async def import_from_docx(docx_files: list[Path], project_id: str, creator_id: str,
                            dry_run: bool = False, extract_only: bool = False):
    """导入 DOCX 文件中的节点。"""
    from emily_core.services.node_commands import (
        CreateNodeCommand, CreateDeliverableCommand, AddDependencyCommand,
    )
    from emily_core.services.node_service import NodeService

    svc = NodeService()
    total_created = 0

    for docx_path in docx_files:
        print(f"\n--- 处理: {docx_path.name} ---")

        text = extract_text_from_docx(str(docx_path))
        print(f"提取文本: {len(text)} 字符")

        if extract_only:
            print(f"\n[提取文本预览]\n{text[:500]}...")
            continue

        nodes = await extract_nodes_with_llm(text)
        if not nodes:
            print("LLM 未提取到节点数据")
            continue

        print(f"LLM 提取到 {len(nodes)} 个节点")

        for nd in nodes:
            if dry_run:
                print(f"  [DRY-RUN] {nd.get('node_id', '?')}: {nd.get('node_name', '?')}")
                total_created += 1
                continue

            try:
                cmd = CreateNodeCommand(
                    project_id=project_id,
                    node_id=nd["node_id"],
                    node_name=nd["node_name"],
                    deadline=nd.get("deadline", ""),
                    owner_dept_id=nd.get("owner_dept_id", "项目总"),
                    stage_id=nd.get("stage_id", 0),
                    creator_id=creator_id,
                    parent_node_id=nd.get("parent_node_id", ""),
                )
                result = await svc.create_node(cmd)
                if result.success:
                    print(f"  [OK] {nd['node_id']}: {nd['node_name']}")
                    total_created += 1

                    for d in nd.get("deliverables", []):
                        await svc.create_deliverable(CreateDeliverableCommand(
                            node_id=nd["node_id"],
                            deliverable_name=d["deliverable_name"],
                            target_amount=d.get("target_amount", 1.0),
                            unit=d.get("unit", "份"),
                            is_required=d.get("is_required", True),
                            operator_id=creator_id,
                        ))

                    for dep in nd.get("dependencies", []):
                        await svc.add_dependency(AddDependencyCommand(
                            node_id=nd["node_id"],
                            depends_on_deliverable_id=dep["depends_on_deliverable_id"],
                            weight=dep.get("weight", 1.0),
                            operator_id=creator_id,
                        ))
                else:
                    print(f"  [FAIL] {nd['node_id']}: {result.message}")
            except Exception as e:
                print(f"  [FAIL] {nd.get('node_id', '?')}: {e}")

    print(f"\n总计导入: {total_created} 个节点")


def main():
    args = parse_args()
    path = Path(args.path)

    if not path.exists():
        print(f"错误：路径不存在: {args.path}")
        sys.exit(1)

    docx_files = []
    if path.is_dir():
        docx_files = sorted(path.glob("*.docx"))
    elif path.suffix.lower() == ".docx":
        docx_files = [path]
    else:
        print(f"错误：不支持的文件格式: {path.suffix}（仅支持 .docx）")
        sys.exit(1)

    if not docx_files:
        print("未找到 DOCX 文件")
        sys.exit(1)

    print(f"找到 {len(docx_files)} 个 DOCX 文件")
    asyncio.run(import_from_docx(
        docx_files, args.project_id, args.creator_id,
        args.dry_run, args.extract_only,
    ))


if __name__ == "__main__":
    main()
