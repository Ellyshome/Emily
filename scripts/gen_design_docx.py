"""生成软著文档鉴别材料 docx —— Emily 企业公共大脑系统设计说明书。

用法: python scripts/gen_design_docx.py
输出: ./Emily_文档鉴别材料.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---- 配置 ----
SOFTWARE_NAME = "Emily 企业公共大脑系统"
VERSION = "V0.7.0"
OUTPUT_FILE = "Emily_文档鉴别材料.docx"
ROOT = Path(__file__).resolve().parent.parent


# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color: str):
    """设置单元格底色。"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(10.5), bold=False, color=None):
    """统一设置 run 的中英文字体。"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)


def add_heading_styled(doc, text, level=1):
    """添加标题段落。"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, cn_font="黑体", en_font="Arial", size=Pt(16 if level == 1 else 14 if level == 2 else 12), bold=True)
    return h


def add_para(doc, text, cn_font="宋体", size=Pt(10.5), bold=False, align=None, indent=True):
    """添加正文段落。"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=bold)
    return p


def add_placeholder_image(doc, caption: str, height_cm: float = 8.0):
    """添加占位图片区域。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # 绘制虚线边框占位区域
    border_text = f"┌{'─'*44}┐\n"
    lines = int(height_cm * 3.5)  # 约3.5行/cm
    for _ in range(lines):
        border_text += f"│{' '*44}│\n"
    border_text += f"└{'─'*44}┘\n"
    border_text += f"  【请替换为: {caption}】"

    run = p.add_run(border_text)
    set_run_font(run, cn_font="楷体", en_font="Courier New", size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    set_run_font(run, cn_font="宋体", size=Pt(9), color=RGBColor(0x40, 0x40, 0x40))


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    """添加格式化表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, cn_font="黑体", size=Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2c3e50")

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, cn_font="宋体", size=Pt(9))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_code_block(doc, code: str, language: str = ""):
    """添加代码块。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)

    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f5f5f5"/>')
    pPr.append(shading)

    run = p.add_run(code)
    set_run_font(run, cn_font="宋体", en_font="Courier New", size=Pt(8))


def add_ascii_diagram(doc, diagram: str, caption: str = ""):
    """添加 ASCII 架构图。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(13)

    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="eef4fa"/>')
    pPr.append(shading)

    run = p.add_run(diagram)
    set_run_font(run, cn_font="宋体", en_font="Courier New", size=Pt(8))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        run = cap.add_run(caption)
        set_run_font(run, cn_font="宋体", size=Pt(9), color=RGBColor(0x40, 0x40, 0x40))


# ============================================================
# 文档主体
# ============================================================

def build_document() -> Document:
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"{SOFTWARE_NAME} {VERSION}  设计说明书"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        set_run_font(run, size=Pt(9), color=RGBColor(0x80, 0x80, 0x80))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("— PAGE —")
    set_run_font(run, size=Pt(9), color=RGBColor(0x80, 0x80, 0x80))

    # ==================== 封面 ====================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SOFTWARE_NAME)
    set_run_font(run, cn_font="黑体", size=Pt(26), bold=True)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("设计说明书")
    set_run_font(run, cn_font="黑体", size=Pt(22), bold=True)

    for _ in range(3):
        doc.add_paragraph()

    for item in [f"软件名称：{SOFTWARE_NAME}", f"版本号：{VERSION}", "文档类别：设计说明书", "编制日期：2026年07月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_run_font(run, cn_font="宋体", size=Pt(14))

    doc.add_page_break()

    # ==================== 目录页 ====================
    add_heading_styled(doc, "目  录", level=1)
    toc_items = [
        ("1", "概述"),
        ("  1.1", "软件概述"),
        ("  1.2", "运行环境"),
        ("  1.3", "设计目标与原则"),
        ("2", "总体设计"),
        ("  2.1", "系统架构设计"),
        ("  2.2", "容器拓扑设计"),
        ("  2.3", "分层架构设计"),
        ("  2.4", "模块总览"),
        ("3", "详细设计"),
        ("  3.1", "消息接入与协议层"),
        ("  3.2", "Session 会话管理"),
        ("  3.3", "WorkItem 任务单元"),
        ("  3.4", "PipelineBUS 四节点管道总线"),
        ("  3.5", "Hook 声明式挂载体系"),
        ("  3.6", "权限与安全体系"),
        ("  3.7", "AI/LLM 集成设计"),
        ("  3.8", "RAG 知识检索设计"),
        ("  3.9", "全景节点图设计"),
        ("  3.10", "系统调度引擎"),
        ("  3.11", "元认知与自进化机制"),
        ("  3.12", "技能系统设计"),
        ("  3.13", "日志与审计体系"),
        ("4", "数据库设计"),
        ("  4.1", "数据库总体架构"),
        ("  4.2", "核心数据表"),
        ("  4.3", "数据关系图"),
        ("5", "接口设计"),
        ("  5.1", "HTTP API 接口"),
        ("  5.2", "SSE 事件推送接口"),
        ("  5.3", "标准协议对象"),
        ("  5.4", "管道阶段接口"),
        ("6", "部署与运维设计"),
        ("7", "功能效果展示"),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        is_main = not num.startswith(" ")
        p.paragraph_format.left_indent = Cm(0) if is_main else Cm(1.0)
        p.paragraph_format.line_spacing = Pt(24)
        run = p.add_run(f"{num}  {title_text}")
        set_run_font(run, cn_font="宋体", size=Pt(12) if is_main else Pt(10.5), bold=is_main)

    doc.add_page_break()

    # ==================== 1. 概述 ====================
    add_heading_styled(doc, "1  概述", level=1)

    add_heading_styled(doc, "1.1  软件概述", level=2)
    add_para(doc, f"{SOFTWARE_NAME} {VERSION}（以下简称 Emily）是一款面向企业的 AI Agent 工具，通过即时通讯平台（QQ）与员工交互，实现团队工作流的数字化记录与留痕、业务 SOP 的智能化引导、企业知识库的 RAG 检索增强生成，以及全景节点图的构建与管理。")
    add_para(doc, "Emily 采用双容器微服务架构：薄插件层部署于 AstrBot 消息平台容器内，负责消息去重（SHA256 指纹）、格式标准化与 HTTP 转发；业务内核部署于独立 FastAPI 容器，承载全部业务逻辑、AI 推理与数据持久化。两个容器通过 HTTP POST 和 SSE（Server-Sent Events）实现双向通信。")
    add_para(doc, "系统核心架构为 Session 主线 + WorkItem + 四节点 PipelineBUS 模型。用户每条消息经 SessionAgent 路由后，可能拆解为一个或多个 WorkItem，每个 WorkItem 在四节点管道总线中顺序执行：意图分析→计划生成→执行验收→成果总结，全程受声明式 Hook 体系管控。")
    add_para(doc, "Emily 的核心创新点包括：(1) 基于声明式 JSON 的 Hook 挂载体系，无需修改框架代码即可扩展安全管控逻辑；(2) SOP 语义匹配驱动的业务流程自动化，SOP 新增仅需放置 Markdown 文件即可生效；(3) 元认知与自进化机制，系统能对自身行为进行反思和改进；(4) 全景节点图的三态流转模型，支持项目进度的精细化追踪。")

    add_heading_styled(doc, "1.2  运行环境", level=2)
    add_para(doc, "系统运行环境要求如下：")
    add_table(doc,
        ["类别", "技术选型", "版本/说明"],
        [
            ["操作系统", "Linux（Docker 容器）", "Ubuntu 22.04 LTS"],
            ["运行时", "Python", "3.10+"],
            ["Web 框架", "FastAPI", "0.100+"],
            ["异步框架", "uvicorn", "ASGI 服务器"],
            ["数据库", "PostgreSQL", "15.x"],
            ["ORM", "SQLAlchemy", "2.0 sync"],
            ["AI/LLM", "DeepSeek API", "OpenAI 兼容协议"],
            ["向量检索", "MaxKB hit_test API", "Qwen3-Embedding-0.6B + pgvector"],
            ["消息平台", "AstrBot + NapCat", "QQ 桥接（OneBot 协议）"],
            ["容器编排", "Docker Compose", "v2"],
            ["包管理", "uv", "Python 高速包管理器"],
            ["邮件集成", "IMAP/SMTP", "企业邮箱收发"],
        ],
        col_widths=[3.5, 4.5, 5.0]
    )

    add_para(doc, "硬件最低配置：CPU 4核、内存 8GB、磁盘 50GB。推荐配置：CPU 8核、内存 16GB、SSD 100GB。数据库服务器需独立磁盘以保证 I/O 性能。")

    add_heading_styled(doc, "1.3  设计目标与原则", level=2)
    add_para(doc, "Emily 的设计目标包括以下六个方面：")
    goals = [
        "工作流留痕：自动识别并记录团队中的事件、任务、会议、文件等业务对象，实现工作过程的可追溯、可查询、可统计。",
        "SOP 数字化：将企业业务流程以 Markdown 文件形式定义为 SOP 手册，通过 LLM 语义匹配驱动流程执行，新增 SOP 无需修改代码。",
        "知识增强：基于企业知识库的 RAG 检索增强生成，为员工提供准确的知识问答服务，支持向量检索和关键词回退双重机制。",
        "全景节点图：构建项目级节点树（含父子权重聚合与三态流转模型），实现项目进度的可视化追踪和自动预警。",
        "安全可控：声明式权限体系与 Hook 管控机制，确保 AI 行为在安全边界内运行，支持行级数据安全策略。",
        "可扩展性：模块化架构设计，新增 SOP/工具/Hook/调度任务无需修改核心框架代码，支持 Mock/Real 模式切换。",
    ]
    for g in goals:
        add_para(doc, f"● {g}")

    add_para(doc, "系统设计遵循以下核心原则：")
    principles = [
        "业务内核独立：emily_core 包不依赖任何 AstrBot 对象，确保核心业务逻辑的可移植性。",
        "分层不可跳：严格遵循 API → EmilyCore → Session → WorkItem → Application → Service → Repository → DB 的调用链路。",
        "SOP 即路由：新增 SOP 等同于新增路由规则，SkillRegistry 管理目录索引，LLM 做语义匹配。",
        "Sync Repo + async Service：Repository 层全部同步，async Service 层通过 asyncio.to_thread() 包裹同步调用。",
        "Hook 三态 deny-wins：ALLOW/WARN/BLOCK 三态决策，任一 BLOCK 即终止管道，确保安全第一。",
    ]
    for pr in principles:
        add_para(doc, f"● {pr}")

    # ==================== 2. 总体设计 ====================
    doc.add_page_break()
    add_heading_styled(doc, "2  总体设计", level=1)

    add_heading_styled(doc, "2.1  系统架构设计", level=2)
    add_para(doc, "Emily 采用分层解耦架构，整体分为五个层次：通信层、协议层、会话层、工作层和基础设施层。各层之间通过标准接口通信，上层依赖下层，不可跨层调用。此设计确保了模块间的低耦合与高内聚，便于独立开发、测试和维护。")

    add_ascii_diagram(doc, """
┌──────────────────────────────────────────────────────────────┐
│                        通信层 (IM)                            │
│     QQ → NapCat → AstrBot → emily_agent 薄插件               │
└─────────────────────────┬────────────────────────────────────┘
                          │ HTTP POST /api/v1/message/send
                          ▼
┌──────────────────────────────────────────────────────────────┐
│                       协议层 (API)                            │
│     FastAPI Routes + SSE Endpoint + Auth Middleware           │
└─────────────────────────┬────────────────────────────────────┘
                          │ EmilyCore.handle_message()
                          ▼
┌──────────────────────────────────────────────────────────────┐
│                      会话层 (Session)                         │
│     DomainTakeover → UserBinding → SessionPool →             │
│     SessionAgent (意图识别 / 快回 / WorkItem 拆分)             │
└─────────────────────────┬────────────────────────────────────┘
                          │ SessionScheduler
                          ▼
┌──────────────────────────────────────────────────────────────┐
│                     工作层 (WorkItem)                         │
│     WorkItem (6态状态机) → PipelineBUS (4节点管道)             │
│     node1:意图分析 → node2:计划生成 →                          │
│     node3:执行验收 → node4:成果总结                            │
│     + Hook 声明式挂载 (Auth/Audit/Trace/Progress)             │
└─────────────────────────┬────────────────────────────────────┘
                          │ Service → Repository
                          ▼
┌──────────────────────────────────────────────────────────────┐
│                   基础设施层 (Infrastructure)                  │
│     PostgreSQL (53表) │ DeepSeek LLM │ MaxKB RAG │           │
│     Email (IMAP/SMTP) │ 文件存储 │ 日志体系                  │
└──────────────────────────────────────────────────────────────┘
""", "图 2-1  Emily 系统分层架构图")

    add_heading_styled(doc, "2.2  容器拓扑设计", level=2)
    add_para(doc, "系统采用 Docker Compose 编排五个容器，各容器职责明确、独立运行，通过 Docker 内部网络互联。薄插件容器（emily_agent）与业务内核容器（emily-core）之间通过 HTTP/SSE 双向通信，实现了消息接入与业务处理的完全解耦。")

    add_ascii_diagram(doc, """
   ┌──────────┐    ┌──────────┐    ┌──────────────────────────┐
   │  NapCat  │───▶│ AstrBot  │───▶│  emily_agent (薄插件)     │
   │  :6099   │    │          │    │  去重+标准化+HTTP转发+SSE  │
   └──────────┘    └──────────┘    └─────────────┬────────────┘
                                                 │ HTTP POST
                                                 ▼
                                   ┌──────────────────────────┐
                                   │   emily-core (FastAPI)    │
                                   │   :18080                  │
                                   │   业务内核 + AI + 全景节点 │
                                   └──────┬──────────┬────────┘
                                          │          │
                           ┌──────────────┘          └──────────────┐
                           ▼                                        ▼
                ┌──────────────────┐                    ┌──────────────────┐
                │ emily-postgres   │                    │     MaxKB        │
                │   :5432          │                    │   :8080 RAG      │
                │   53表 PostgreSQL │                    │ Qwen3-Embedding  │
                └──────────────────┘                    └──────────────────┘
""", "图 2-2  容器拓扑图")

    add_table(doc,
        ["容器", "端口", "说明", "技术栈"],
        [
            ["napcat", "6099 (WebUI)", "QQ 桥接服务", "NapCat + OneBot协议"],
            ["astrbot", "—", "消息平台 + 薄插件宿主", "AstrBot + Python"],
            ["emily-core", "18080", "业务内核（FastAPI）", "FastAPI + emily_core包"],
            ["maxkb", "8080", "知识库 RAG 服务", "MaxKB + pgvector"],
            ["emily-postgres", "5432", "独立 PostgreSQL 数据库", "PostgreSQL 15"],
        ],
        col_widths=[2.5, 2.5, 4.0, 4.0]
    )

    add_heading_styled(doc, "2.3  分层架构设计", level=2)
    add_para(doc, "系统严格遵循分层调用原则，各层职责和依赖关系如下。上层只能调用相邻下层的接口，禁止跨层调用。此设计确保了模块间的低耦合与高内聚，便于独立开发、测试和维护。调用链路为：API层 → EmilyCore → Session → WorkItem → Application → Service → Repository → DB。")

    add_table(doc,
        ["层次", "包路径", "核心职责", "关键类/模块"],
        [
            ["协议层 (API)", "emily-core/api/", "FastAPI 路由、SSE 端点、认证中间件", "server.py, routes/, sse/"],
            ["内核编排 (Core)", "emily_core/", "统一入口、系统初始化、子系统编排", "EmilyCore, bootstrap.py"],
            ["适配层 (Adapter)", "emily_core/adapters/", "消息标准化、Session 工厂与池管理", "StandardMessage, SessionPool"],
            ["会话层 (Session)", "emily_core/session/", "意图识别、快回通道、WorkItem 拆分", "SessionAgent, FocusLock"],
            ["工作层 (WorkItem)", "emily_core/workitem/", "WorkItem 状态机、PipelineBUS、Hook", "WorkItem, PipelineBUS"],
            ["应用层 (Application)", "emily_core/application/", "事件/任务/会议/文件/节点/权限/查询", "EventApp, TaskApp 等"],
            ["服务层 (Service)", "emily_core/services/", "业务逻辑实现、30+核心服务", "NodeService, PermissionService"],
            ["仓库层 (Repository)", "emily_core/repositories/", "数据访问封装、SQLAlchemy ORM", "19个Repo类"],
            ["基础设施层", "emily_core/infrastructure/", "数据库(53表)、LLM、日志体系", "models.py, client.py"],
            ["工具层 (Tools)", "emily_core/tools/", "业务流工具、SOP 工具注册与执行", "BusinessFlowTool, 15+工具"],
            ["权限层 (Permission)", "emily_core/permission/", "权限引擎、行级安全、缓存", "AuthEngine, RowSecurity"],
            ["调度层 (Scheduler)", "emily_core/scheduler/", "调度引擎、Job 注册、周期性任务", "SchedulerEngine"],
            ["技能层 (Skill)", "emily_core/skill/", "Skill 定义、解析、参数提取、执行器", "SkillRegistry, Executor"],
        ],
        col_widths=[2.5, 3.5, 4.0, 3.0]
    )

    add_heading_styled(doc, "2.4  模块总览", level=2)
    add_para(doc, "系统包含约 200 个 Python 源文件，总代码量约 35,000 行，按功能划分为以下主要模块：")

    modules = [
        ("消息处理", "emily_core/session/", "SessionAgent、SessionPool、FocusLock、ConfirmQueue、数据获取器（5个 fetcher）"),
        ("工作流引擎", "emily_core/workitem/", "WorkItem 状态机、WorkItemAgent、PipelineBUS 四节点管道、Hook 体系、BusContext"),
        ("应用服务", "emily_core/application/", "EventApp、TaskApp、MeetingApp、FileApp、NodeApp、PermissionApp、QueryApp 共7个应用"),
        ("核心服务", "emily_core/services/", "30+ 业务服务，覆盖权限、节点、进化、RAG、查询、邮件、日志等领域"),
        ("数据访问", "emily_core/repositories/", "19 个仓库类，封装 User/Node/Event/Task/Message/Scheduler 等53表的CRUD"),
        ("权限安全", "emily_core/permission/", "AuthEngine（权限引擎）、RowSecurity（行级安全）、CodeCompiler（权限码编译）、PermissionCache"),
        ("AI 集成", "emily_core/infrastructure/llm/", "LLM 客户端（chat/chat_json/chat_with_tools 三种模式）、Prompt 加载器"),
        ("数据库", "emily_core/infrastructure/database/", "53 表 ORM 模型、Session 工厂、Schema 自动补齐"),
        ("调度系统", "emily_core/scheduler/", "SchedulerEngine（tick循环+Advisory Lock）、HandlerRegistry、11 个 Job 实现"),
        ("技能系统", "emily_core/skill/", "SkillRegistry（技能注册）、SkillExecutor（执行器）、ParamExtractor、Parser、Validator"),
        ("工具注册", "emily_core/tools/", "15+ 工具（事件/任务/节点/查询/邮件/知识搜索/文件/记忆/待办/归档等）"),
        ("RAG 检索", "emily_core/providers/rag/", "MaxKB Provider（主检索）、本地关键词回退（local_fallback）"),
        ("日志体系", "emily_core/infrastructure/logging/", "9 个专用 Logger：业务事件/反馈/LLM/管道/RAG/调度/会话生命周期等"),
    ]
    for name, path, desc in modules:
        add_para(doc, f"● {name}（{path}）：{desc}")

    # ==================== 3. 详细设计 ====================
    doc.add_page_break()
    add_heading_styled(doc, "3  详细设计", level=1)

    # 3.1
    add_heading_styled(doc, "3.1  消息接入与协议层", level=2)
    add_para(doc, "消息接入采用薄插件模式。emily_agent 作为 AstrBot 插件运行，仅负责消息去重（SHA256 指纹）、格式标准化、HTTP 转发和 SSE 监听，不包含任何业务逻辑。所有业务处理在独立 emily-core 容器中完成。此设计使得业务内核与消息平台完全解耦，可以方便地移植到其他 IM 平台。")

    add_ascii_diagram(doc, """
  IM 消息到达                     薄插件处理                       Core 处理
  ──────────                     ──────────                       ─────────
  QQ 消息 ──▶ NapCat ──▶ AstrBot ──▶ emily_agent
                                            │
                                       ┌────┴────┐
                                       │ 1.SHA256 │ 消息去重（防重复处理）
                                       │   去重   │
                                       │ 2.标准化 │ 转为 StandardMessage
                                       │   转换   │
                                       │ 3.HTTP   │ POST /api/v1/message/send
                                       │   转发   │
                                       └────┬────┘
                                            │
                                            ▼
                                      EmilyCore.handle_message()
                                            │
                                       ┌────┴────┐
                                       │ Domain   │ 接管判断（是否处理此消息）
                                       │ Takeover │
                                       │ User     │ IM→系统用户自动绑定
                                       │ Binding  │
                                       └────┬────┘
                                            │
                                            ▼
                                       SessionPool.route()
""", "图 3-1  消息接入与处理流程图")

    add_para(doc, "标准协议对象是跨容器通信的核心合约，定义在 emily_core/adapters/standard/ 目录下。入站消息统一为 StandardMessage，包含消息 ID、平台标识、会话类型（私聊/群聊）、发送者信息、消息内容、附件等字段。出站回复统一为 ReplyMessage，支持文本内容、引用回复、知识来源引用和文件附件。路由决策 RouteDecision 携带接管判断结果和置信度。")
    add_para(doc, "此外还定义了 RouteResult（结构化参数载体，携带意图类型和业务数据）、HandlerResult（Application 层输出，包含成功标志和业务对象 ID）、AgentStep/AgentResult（ReAct 循环输出）等辅助对象。Command DTO（EventCommand、TaskCommand、MeetingCommand、FileCommand、QueryCommand）作为 Application→Service 的传参信封，确保了层间数据传递的类型安全。")

    # 3.2
    add_heading_styled(doc, "3.2  Session 会话管理", level=2)
    add_para(doc, "SessionAgent 是每个会话（群聊或私聊）的智能大脑，由 SessionPoolManager 按 conversation_id 路由和维护生命周期。SessionPoolManager 内部维护 conversation_id → SessionAgent 的映射表，新会话自动创建 SessionAgent，过期会话自动清理。")

    add_ascii_diagram(doc, """
  SessionAgent.handle(StandardMessage)
       │
       ├─① 快回通道 ──▶ 问候/感谢/告别 ──▶ 直接友好回复（无LLM，秒级响应）
       │
       └─② 意图识别 ──▶ LLM + SOP 知识库 ──▶ SkillRegistry 语义匹配
              │
              ├─ 单 SOP 匹配 ──▶ 创建 1 个 WorkItem (sop_id=命中SOP)
              │
              ├─ 复合请求   ──▶ 拆解为 N 个 WorkItem（支持依赖与优先级）
              │
              └─ 未命中     ──▶ 创建 1 个 WorkItem (sop_id=__FALLBACK_SOP__)
              │
              ▼
       SessionScheduler ──▶ 排队 + 状态机驱动 ──▶ PipelineBUS
""", "图 3-2  SessionAgent 处理流程图")

    add_para(doc, "SessionAgent 内部维护两个关键组件：FocusLock（话题切换检测，防止上下文漂移）和 ConfirmQueue（待确认队列，管理需要用户确认的操作）。SessionContext 对象封装了会话级的所有上下文信息，包括当前用户、项目、RAG 缓存、可见文件列表等。")
    add_para(doc, "数据获取阶段由 5 个专用 Fetcher 完成：fetch_available_tools（获取可用工具列表）、fetch_rag_info（RAG 知识检索）、fetch_system_description（系统描述注入）、fetch_visible_files（可见文件列表）、fetch_visible_schema（可见数据 Schema）。这些 Fetcher 的结果注入到 SessionAgent 的上下文中，为后续意图识别和 LLM 调用提供丰富的环境信息。")

    # 3.3
    add_heading_styled(doc, "3.3  WorkItem 任务单元", level=2)
    add_para(doc, "WorkItem 是系统中最小的独立执行单元，采用六态状态机管理生命周期。一条用户消息可能产生零到多个 WorkItem，每个 WorkItem 在 PipelineBUS 中独立流转。WorkItem 对象封装了完整的执行上下文：SOP ID、路由决策、执行计划、步骤结果、最终回复等，实现了一个 WorkItem 即一个任务全息记录。")

    add_ascii_diagram(doc, """
  ┌──────────┐     ┌──────────┐     ┌───────────┐     ┌──────┐
  │ CREATED  │────▶│ PLANNING │────▶│ EXECUTING │────▶│ DONE │
  └──────────┘     └──────────┘     └───────────┘     └──────┘
       │                │                │
       │                │                │         ┌────────────────┐
       │                │                └────────▶│ WAITING_CONFIRM│
       │                │                           └───────┬────────┘
       │                │                                   │
       ▼                ▼                                   ▼
  ┌──────────┐     ┌──────────┐                      ┌───────────┐
  │  FAILED  │◀────│  FAILED  │◀─────────────────────│  EXECUTING │
  └──────────┘     └──────────┘                      └───────────┘
""", "图 3-3  WorkItem 六态状态机")

    add_table(doc,
        ["状态", "含义", "驱动者", "说明"],
        [
            ["CREATED", "刚创建，未开始执行", "SessionScheduler", "WorkItem 初始状态"],
            ["PLANNING", "正在生成执行计划", "SessionScheduler._run_one()", "PipelineBUS node2 生成 ExecutionPlan"],
            ["EXECUTING", "正在执行计划步骤", "PipelineBUS node3", "遍历 PlanStep，调用工具"],
            ["WAITING_CONFIRM", "等待用户确认后继续", "预留状态", "已定义，当前未驱动进入此状态"],
            ["DONE", "执行成功完成（终态）", "PipelineBUS node4 后", "成果已总结，回复已发送"],
            ["FAILED", "执行异常或终止（终态）", "PipelineBUS 异常处理", "Hook BLOCK 或执行异常"],
        ],
        col_widths=[2.5, 3.0, 3.5, 5.0]
    )

    # 3.4
    add_heading_styled(doc, "3.4  PipelineBUS 四节点管道总线", level=2)
    add_para(doc, "PipelineBUS 是 WorkItem 的执行引擎，采用固定顺序的四节点管道模型。消息必须按 node1→node4 顺序流经每个节点，节点之间通过 BusContext 传递上下文。每个节点执行前后均可触发声明式 Hook，实现安全管控和审计追踪。")

    add_ascii_diagram(doc, """
  ┌────────────────────────────────────────────────────────────────────────┐
  │                        PipelineBUS 四节点管道                           │
  │                                                                        │
  │  ┌─────────────┐   ┌─────────────┐   ┌─────────────┐   ┌────────────┐│
  │  │   node1     │   │   node2     │   │   node3     │   │   node4    ││
  │  │  意图+拆分   │──▶│  计划+标准   │──▶│  执行+验收   │──▶│  成果总结  ││
  │  │  (required) │   │  (required) │   │  (required) │   │  (non-req) ││
  │  └─────────────┘   └─────────────┘   └─────────────┘   └────────────┘│
  │                                                                        │
  │  Hook 触发点:                                                          │
  │  before:node2 ──▶ auth.admin_check + plan_task_match                  │
  │  before:node3 ──▶ trace.reasoning_start + auth.admin_check            │
  │  after:node3  ──▶ trace.reasoning_end + audit.sop_completed          │
  │  before:node4 ──▶ guardian.reply_review + guardian.deep_audit         │
  │  after:node4  ──▶ audit.outbound_archived                            │
  │  on_error:node3 ──▶ audit.agent_error                                 │
  └────────────────────────────────────────────────────────────────────────┘
""", "图 3-4  PipelineBUS 四节点管道与 Hook 挂载点")

    add_table(doc,
        ["节点", "名称", "必选", "核心职责", "Mock 实现", "Real 实现"],
        [
            ["node1", "意图+拆分", "是", "增量灌注 SOP/工具/Schema，构建 RouteDecision", "MockPlanning", "KnowledgeInjector"],
            ["node2", "计划+标准", "是", "生成 ExecutionPlan（risk_level/steps/criteria）", "MockPlanner", "LLMPlanner"],
            ["node3", "执行+验收", "是", "遍历 PlanStep，调用 BusinessFlowTool 或 RAG", "MockWorkAgent", "RealExecutor"],
            ["node4", "成果总结", "否", "合成 result_text → verified_reply", "—", "结果合成器"],
        ],
        col_widths=[1.5, 2.0, 1.5, 4.5, 2.5, 3.0]
    )

    add_para(doc, "BusContext 是管道上下文对象，在节点间传递完整的执行环境：WorkItem、RouteDecision、ExecutionPlan、StepResult 列表等。每个节点通过 BusContext 读取上游节点的输出，写入本节点的产出。此设计使得节点间完全解耦，任何节点都可以独立替换为 Mock 或 Real 实现。")
    add_para(doc, "ExecutionPlan 由 PlanStep 数组组成，每个 PlanStep 包含步骤 ID、描述、工具名称、工具参数、预期输出和依赖关系。RiskLevel 分为三级：L1（低风险，如闲聊和读取操作）、L2（中风险，如写入操作）、L3（高风险，如复合请求和删除操作）。风险等级影响 Hook 的管控策略。")

    # 3.5
    add_heading_styled(doc, "3.5  Hook 声明式挂载体系", level=2)
    add_para(doc, "Hook 体系采用声明式 JSON 配置 + 实现类注册的模式。新增 Hook 只需编辑 hook_config.json 文件并添加实现类，无需修改核心框架代码。Hook 实现类继承自 Hook 基类（emily_core/workitem/pipeline/hook.py），提供 execute() 方法和决策输出。")

    add_para(doc, "Hook 执行遵循三条核心规则：")
    rules = [
        "三态决策：ALLOW（放行）、WARN（警告继续）、BLOCK（终止管道）。",
        "deny always wins：多个 Hook 中任一返回 BLOCK，管道立即终止。",
        "异常即阻断：before Hook 抛异常视为 BLOCK（安全第一原则）；after Hook 异常不阻断，仅记录日志。",
    ]
    for i, r in enumerate(rules, 1):
        add_para(doc, f"({i}) {r}")

    add_para(doc, "当前已注册四种 Hook 类型，共 10 个挂载点：")

    add_table(doc,
        ["挂载点", "Hook 名称", "类型", "决策", "说明"],
        [
            ["before:wi_node2", "auth.admin_check", "auth", "ALLOW/BLOCK", "管理员鉴权，验证用户权限等级"],
            ["before:wi_node2", "plan_task_match", "plan_task_match", "ALLOW", "计划任务周期匹配"],
            ["before:wi_node3", "trace.reasoning_start", "trace", "ALLOW", "创建 Agent 推理追踪记录"],
            ["before:wi_node3", "auth.admin_check", "auth", "ALLOW/BLOCK", "管理员鉴权（执行阶段）"],
            ["after:wi_node3", "trace.reasoning_end", "trace", "ALLOW", "更新推理追踪结果"],
            ["after:wi_node3", "audit.sop_completed", "audit", "ALLOW", "审计记录 SOP 执行完成"],
            ["on_error:wi_node3", "audit.agent_error", "audit", "ALLOW", "审计记录 Agent 执行失败"],
            ["before:wi_node4", "guardian.reply_review", "verify", "ALLOW/WARN", "核验回复内容合规性"],
            ["before:wi_node4", "guardian.deep_audit", "deep_audit", "ALLOW", "深度审计调查（默认关闭）"],
            ["after:wi_node4", "audit.outbound_archived", "audit", "ALLOW", "审计记录回复已发送"],
        ],
        col_widths=[2.5, 3.0, 2.0, 2.0, 4.5]
    )

    add_code_block(doc, '// hook_config.json 配置示例\n{\n  "hooks": [\n    {\n      "name": "auth.admin_check",\n      "mount_point": "before:wi_node2",\n      "type": "auth",\n      "enabled": true,\n      "config": {"admin_only": false}\n    },\n    {\n      "name": "guardian.reply_review",\n      "mount_point": "before:wi_node4",\n      "type": "verify",\n      "enabled": true,\n      "config": {"strict_mode": false}\n    }\n  ]\n}')

    # 3.6
    add_heading_styled(doc, "3.6  权限与安全体系", level=2)
    add_para(doc, "权限体系包含四个核心组件，共同构建了从请求级到数据行的全链路安全管控：")
    add_para(doc, "(1) AuthEngine（权限引擎）：基于 SOP allow_roles 匹配用户角色，输出 ALLOW/DENY 决策。读取 User 的 perm_list 和 grouping 信息，与 SOP 定义的允许角色列表进行匹配。")
    add_para(doc, "(2) RowSecurity（行级安全策略）：控制数据的可见范围，基于 PublicFieldRegistry（公开字段白名单）和 PermissionDef（权限码定义）实现字段级访问控制。不同权限等级的用户只能访问对应等级的数据字段。")
    add_para(doc, "(3) CodeCompiler（权限码编译器）：将资源路径（resource_type/security_level/project_id/node_id/resource-id）编译为权限码，用于精确的权限匹配和校验。")
    add_para(doc, "(4) PermissionCache（权限缓存）：缓存用户的权限信息，减少重复的数据库查询，提高鉴权性能。缓存通过用户 ID 和权限组进行索引，支持按需刷新。")
    add_para(doc, "权限授权支持三种类型：AUTO（系统自动授予，如项目成员自动获得项目相关权限）、TEMP（临时授权，有有效期，到期自动失效）、PERMANENT（永久授权，需管理员手动撤销）。所有权限变更通过 PermissionAuditLog 记录，该表仅允许 INSERT，不可篡改，确保审计完整性。")
    add_para(doc, "API 层通过 AuthMiddleware 中间件实现请求级鉴权，验证 API Key 和来源合法性。SOP 层面，每个 SOP 手册定义 allow_roles 字段，限定哪些角色可以触发该 SOP。")

    # 3.7
    add_heading_styled(doc, "3.7  AI/LLM 集成设计", level=2)
    add_para(doc, "系统通过 LLM 客户端（emily_core/infrastructure/llm/client.py）与 DeepSeek API 交互，支持三种调用模式：")

    add_table(doc,
        ["模式", "方法", "用途", "返回格式", "Token消耗"],
        [
            ["普通对话", "chat()", "闲聊快回、意图识别", "自然语言文本", "低"],
            ["结构化输出", "chat_json()", "SOP 匹配、参数提取", "JSON 对象", "中"],
            ["工具调用", "chat_with_tools()", "BusinessFlowTool 参数生成", "function_call 结果", "高"],
        ],
        col_widths=[2.0, 2.5, 3.5, 2.5, 2.0]
    )

    add_para(doc, "LLM 调用全过程通过 LlmInteractionLog 记录，包含模型名称、输入/输出 Token 数、响应延迟、响应类型等指标，用于性能监控和成本分析。Prompt 模板通过 PromptLoader 从 emily-data/prompts/ 目录加载，支持 session、workitem、guardian_step、guardian_reply、project 等多套 Prompt 模板，可根据不同场景灵活切换。")
    add_para(doc, "在 M14 结构化输出优先模式下，命中 SOP 后 LLM 通过 chat_json() 输出 {tool, params} 结构，框架直接调用 BusinessFlowTool.handler(params) 执行业务逻辑，避免了 LLM function-calling 的不可控风险。未命中 SOP 时，通过 SkillExecutor 兜底处理。此设计在保证执行确定性的同时，充分利用了 LLM 的语义理解能力。")

    # 3.8
    add_heading_styled(doc, "3.8  RAG 知识检索设计", level=2)
    add_para(doc, "RAG（检索增强生成）模块采用双引擎架构：MaxKB hit_test API 作为主检索引擎，本地关键词匹配作为回退方案。主检索使用 Qwen3-Embedding-0.6B 向量模型将查询文本编码为向量，通过 pgvector 在 PostgreSQL 中进行近似最近邻搜索，返回相关度最高的知识片段。")

    add_ascii_diagram(doc, """
  用户查询 ──▶ RAG Provider ──┬──▶ MaxKB hit_test API ──▶ 向量检索结果
                              │    (Qwen3-Embedding-0.6B)
                              │    (pgvector ANN Search)
                              │
                              └──▶ 本地关键词回退 ──▶ TF-IDF 匹配结果
                                   (MaxKB 不可用时激活)
""", "图 3-5  RAG 双引擎检索架构")

    add_para(doc, "RAG 检索在两个阶段被调用：一是在 SessionAgent 的数据获取阶段（fetch_rag_info），为意图识别提供知识上下文；二是在 PipelineBUS 的 node3 执行阶段（knowledge_search），为业务流程提供知识支撑。检索结果通过 RagResult 对象封装，包含来源文档、相关度分数和检索内容，并可通过 references 字段实现知识溯源。")

    # 3.9
    add_heading_styled(doc, "3.9  全景节点图设计", level=2)
    add_para(doc, "全景节点图是 Emily 的项目管理核心功能，采用树状结构组织项目中的工作节点。节点类型分为四级：MILESTONE（里程碑）、STAGE（阶段）、OPERATION（工序）、DELIVERABLE（交付物）。每个节点支持父子关系、权重聚合和三态流转模型（ACTIVE/COMPLETED/DISCARDED）。")

    add_ascii_diagram(doc, """
  ┌───────────┐
  │ MILESTONE │  里程碑（如：主体结构封顶）
  └─────┬─────┘
        │
   ┌────┴────┐
   ▼         ▼
┌──────┐ ┌──────┐
│STAGE │ │STAGE │  阶段（如：样板段、标准段）
└──┬───┘ └──┬───┘
   │        │
   ▼        ▼
┌────────┐ ┌────────┐
│OPERATION│ │OPERATION│  工序（如：放线、绑筋、浇筑）
└───┬────┘ └───┬────┘
    │          │
    ▼          ▼
┌──────────┐ ┌──────────┐
│DELIVERABLE│ │DELIVERABLE│  交付物（如：施工方案、验收记录）
└──────────┘ └──────────┘
""", "图 3-6  全景节点树层级结构")

    add_para(doc, "NodeStateMachine 管理节点的生命周期流转，实现状态变更的合法性校验。批量操作通过 NodeBatch 服务（create_node_tree）实现，支持从 YAML 配置一键创建完整节点树。节点进度更新支持文本和语音两种输入方式（NodeVoiceEntryTool）。NodeService（1,324 行代码）是系统中最大的服务类，提供完整的节点 CRUD、状态变更、进度追踪、批量操作等功能。")

    # 3.10
    add_heading_styled(doc, "3.10  系统调度引擎", level=2)
    add_para(doc, "SchedulerEngine 是 Emily 的定时任务调度引擎，采用 tick 循环 + PostgreSQL Advisory Lock 实现分布式安全调度。引擎每秒执行一次 tick，检查是否有到期需要执行的作业。Advisory Lock 确保在多实例部署时只有一个实例执行调度任务，避免重复执行。")

    add_para(doc, "调度类型支持三种：ONCE（一次性，指定时间执行一次）、CRON（Cron 表达式，按周期重复执行）、INTERVAL（固定间隔，每隔 N 秒/分/时执行一次）。作业通过 HandlerRegistry 插件式注册，新增作业只需实现 JobHandler 接口并注册到 Registry，无需修改调度引擎代码。")

    add_table(doc,
        ["作业名称", "功能说明", "调度类型"],
        [
            ["periodic_node", "定期创建 TASK 节点（替代旧 PlanTask 循环模板）", "CRON"],
            ["node_deadlines", "节点截止日期检查与提前预警提醒", "CRON"],
            ["morning_report", "晨报生成与推送（每日摘要）", "CRON"],
            ["daily_insight", "每日洞察分析（发现异常和模式）", "CRON"],
            ["rule_induction", "规则归纳与学习（从历史交互中提炼规则）", "INTERVAL"],
            ["patch_validator", "自进化补丁验证（确保补丁安全性）", "INTERVAL"],
            ["system_description_update", "系统描述自动更新（元认知）", "INTERVAL"],
            ["world_book_update", "世界知识库更新（同步企业数据）", "INTERVAL"],
            ["session_cleanup", "过期会话清理（释放资源）", "CRON"],
            ["health_check", "系统健康检查（数据库/LLM/RAG连通性）", "INTERVAL"],
            ["data_sync", "数据同步（与外部系统对账）", "CRON"],
        ],
        col_widths=[3.5, 6.5, 2.5]
    )

    add_para(doc, "每次调度执行的结果记录在 scheduler_executions 表中，包含执行状态（PENDING/RUNNING/SUCCESS/FAILED）、开始时间、结束时间和错误信息。SchedulerService 提供作业的 CRUD 操作和执行记录查询。")

    # 3.11
    add_heading_styled(doc, "3.11  元认知与自进化机制", level=2)
    add_para(doc, "Emily 具备元认知能力，能够对自身行为进行反思和改进。该机制包含以下核心组件：")

    components = [
        "SystemDescriptionBuilder：自动构建系统描述，汇总数据库 Schema、SOP 清单、工具列表等信息，为 AI 提供自省上下文。这是系统中代码量最大的服务之一（537行）。",
        "SystemDescriptionService：管理系统描述的版本化和增量更新，配合调度作业实现定期自动刷新。",
        "SchemaDriftDetector：检测数据库 Schema 与代码 ORM 模型之间的偏差，发现新增/缺失的列和表，确保数据一致性。",
        "CognitionDriftDetector：检测 AI 认知偏差，通过对比历史回复质量和当前表现，发现退化趋势并预警。",
        "PatchGenerator / PatchApplier / PatchValidator：自进化补丁的三步闭环。PatchGenerator 根据偏差生成修复补丁，PatchValidator 在沙箱中验证补丁安全性，PatchApplier 将通过验证的补丁应用到系统。所有补丁应用记录完整留痕。",
        "RuleInductor：规则归纳器，从历史交互中学习并提炼可复用的行为规则，持续优化系统行为。",
        "InsightGenerator：洞察生成器，从数据中发现模式和异常，为晨报和每日洞察提供素材。",
    ]
    for c in components:
        add_para(doc, f"● {c}")

    add_para(doc, "自进化遵循安全第一原则：所有补丁必须经过 PatchValidator 验证后才能应用，且应用记录完整留痕。PatchValidator 会在隔离环境中执行补丁，验证其不会破坏现有功能，只有全部验证通过才会应用到生产环境。")

    # 3.12
    add_heading_styled(doc, "3.12  技能系统设计", level=2)
    add_para(doc, "技能系统（Skill）是 Emily 的可扩展执行框架，当消息未命中任何 SOP 时，SkillExecutor 兜底处理。技能系统包含以下组件：")

    add_para(doc, "(1) SkillRegistry：技能注册中心，管理技能的发现、索引和匹配。技能以 Markdown 文件形式定义在 emily-data/skills/ 目录下，重启后自动加载。")
    add_para(doc, "(2) SkillExecutor：技能执行器，协调技能的解析、参数提取和调用。执行过程遵循 ReAct（Reason-Act）循环，支持多步推理。")
    add_para(doc, "(3) ParamExtractor：参数提取器，从用户消息中提取技能所需的参数，支持必选/可选参数和类型转换。")
    add_para(doc, "(4) SkillParser：技能解析器，将 Markdown 格式的技能定义文件解析为可执行的结构化对象。")
    add_para(doc, "(5) SkillValidator：技能验证器，验证技能定义的完整性和合法性，确保参数类型、必选项等约束满足。")

    # 3.13
    add_heading_styled(doc, "3.13  日志与审计体系", level=2)
    add_para(doc, "系统实现了 9 个专用 Logger，覆盖业务事件的完整生命周期：")

    add_table(doc,
        ["Logger", "职责", "输出目标"],
        [
            ["BusinessEventLogger", "业务事件日志（创建/更新/删除）", "控制台 + 数据库"],
            ["FeedbackDetector", "用户反馈检测（正面/负面/中性）", "控制台"],
            ["LlmLogger", "LLM 调用详情（模型/token/延迟）", "llm_interaction_logs 表"],
            ["PipelineLogger", "管道节点执行日志", "hook_execution_logs 表"],
            ["RagLogger", "RAG 检索日志（查询/结果/耗时）", "控制台"],
            ["SchedulerLogger", "调度任务执行日志", "scheduler_executions 表"],
            ["SessionLifecycleLogger", "会话创建/销毁/超时日志", "控制台"],
            ["LegacyLogBridge", "旧日志格式桥接（兼容性）", "控制台"],
            ["LogWriter", "统一日志写入器", "文件 + 控制台"],
        ],
        col_widths=[3.5, 5.5, 3.5]
    )

    add_para(doc, "审计数据通过以下四张专用表持久化：agent_reasoning_logs（Agent 推理记录）、llm_interaction_logs（LLM 交互日志）、tool_call_logs（工具调用日志）、hook_execution_logs（Hook 执行日志）。这四张表共同构成了 AI 行为的完整审计链，支持事后追溯和问题定位。")


    # 3.14 SOP 业务流系统设计
    add_heading_styled(doc, "3.14  SOP 业务流系统设计", level=2)
    add_para(doc, "SOP（标准操作流程）是 Emily 实现业务流程数字化的核心机制。SOP 以 Markdown 文件形式定义，放置于 emily-data/sops/ 目录下。每个 SOP 文件包含三个核心部分：流程描述（触发条件、适用范围）、工具定义（结构化输出 Schema）和执行步骤（分步骤操作指引）。")
    add_para(doc, "SOP 的生命周期管理包括：注册（SkillRegistry 扫描目录并索引）、匹配（LLM 语义匹配用户意图与 SOP）、执行（PipelineBUS 中按步骤执行）和审计（SOP 执行过程完整记录于 sop_routing_logs 和 agent_reasoning_logs）。当前系统已注册 10 份 SOP 手册，覆盖事件记录、任务管理、文件归档、会议记录、设计变更等核心业务场景。")
    add_para(doc, "SOP 匹配采用语义相似度算法，LLM 从 SOPIntentRegistry.dump_type_tree() 获取完整的 SOP 知识树，注入 system prompt 后由 LLM 判断用户意图最匹配的 SOP。匹配结果分为三种：单 SOP 匹配（高置信度，直接执行）、复合匹配（用户消息涉及多个 SOP，拆分为多个 WorkItem 依次执行）和未命中（走 SkillExecutor 兜底处理）。")

    add_ascii_diagram(doc, """
  SOP 生命周期：
  
  编写 .md ──▶ 放入 sops/ ──▶ SkillRegistry 自动索引
                                       │
  用户消息 ──▶ LLM 语义匹配 ──▶ 命中 SOP
                                       │
                    ┌──────────────────┼──────────────────┐
                    ▼                  ▼                  ▼
              单 SOP 匹配        复合匹配           未命中
              创建1个WorkItem    拆分N个WorkItem   SkillExecutor兜底
                    │                  │                  │
                    └──────────────────┼──────────────────┘
                                       ▼
                              PipelineBUS 执行
                                       │
                                       ▼
                              sop_routing_logs 审计记录
""", "图 3-7  SOP 业务流生命周期")

    # 3.15 工具注册与执行框架
    add_heading_styled(doc, "3.15  工具注册与执行框架", level=2)
    add_para(doc, "Emily 采用双层工具架构：BusinessFlowTool（业务流工具，M14 框架直调）和 ToolDefinition（LLM 函数调用工具，OpenAI 兼容格式）。两层工具分别服务于不同的调用场景，确保执行确定性。")
    add_para(doc, "BusinessFlowTool 是 SOP 执行阶段的核心工具载体。当 LLM 在 node2 生成 ExecutionPlan 后，node3 中的 RealExecutor 直接调用 BusinessFlowTool.handler(params) 执行业务逻辑，不经过 LLM function-calling。此设计避免了 LLM 工具调用的不确定性，确保业务操作的精确性和可重复性。每个 BusinessFlowTool 的 name 与 SOP 手册中的工具名称一一对应。")
    add_para(doc, "ToolDefinition 则用于 LLM function-calling 场景，导出为 OpenAI function-calling 格式供 LLM 选择调用。ToolRegistry 管理所有工具的注册、索引和查询，支持 register()、get()、list_all()、list_public() 和 get_openai_tools() 等操作。管理员专用工具（require_admin=True）不会出现在普通用户的工具列表中。")

    add_table(doc,
        ["工具名称", "所属类别", "功能说明"],
        [
            ["EventTool", "业务流工具", "事件创建/更新/查询"],
            ["TaskTool", "业务流工具", "任务创建/更新/分配/查询"],
            ["MeetingTool", "业务流工具", "会议记录创建/更新/查询"],
            ["FileTool", "业务流工具", "文件上传/下载/版本管理"],
            ["NodeTool", "全景节点工具", "节点树创建/查询/更新/进度管理"],
            ["NodeTaskTool", "全景节点工具", "节点关联任务创建与管理"],
            ["NodeVoiceEntryTool", "全景节点工具", "语音方式录入节点进度"],
            ["QueryTool", "查询工具", "多维度数据查询（事件/任务/会议/文件等）"],
            ["KnowledgeSearchTool", "RAG 工具", "企业知识库检索"],
            ["EmailTool", "外部集成工具", "邮件发送（SMTP）"],
            ["MemoryTool", "用户工具", "用户记忆存取（偏好/历史摘要）"],
            ["PendingIssueTool", "待办工具", "待办事项创建与追踪"],
            ["ChatArchiveTool", "归档工具", "对话记录归档与检索"],
        ],
        col_widths=[3.5, 3.0, 6.5]
    )

    # 3.16 邮件集成设计
    add_heading_styled(doc, "3.16  邮件集成设计", level=2)
    add_para(doc, "Emily 通过 EmailService 集成企业邮件系统，支持 IMAP 接收和 SMTP 发送。邮件集成主要用于两个场景：一是业务通知邮件的自动发送（如任务分配通知、截止日期提醒等），二是通过邮件桥接扩展 Emily 的触达渠道。")
    add_para(doc, "IMAP Provider（imap_provider.py）负责邮件接收，支持按文件夹、未读标记和时间范围筛选邮件。SMTP Provider（smtp_provider.py）负责邮件发送，支持 HTML 格式邮件和附件。邮件服务通过 EmailTool 暴露给工具执行框架，可在 SOP 流程中被调用。所有邮件操作通过 EmailCommand DTO 传递参数，确保接口类型安全。")

    # 3.17 世界知识库设计
    add_heading_styled(doc, "3.17  世界知识库设计", level=2)
    add_para(doc, "世界知识库（World Book）是 Emily 的外部知识注入机制，用于将企业特定的行业知识、法规标准、项目规范等信息注入到 AI 的上下文中。WorldBookBuilder（568行代码）负责知识库的构建和更新，WorldBookService 提供知识条目的 CRUD 管理。")
    add_para(doc, "知识条目按项目和领域分类存储在 world_book 表中，每条记录包含关键词、触发条件和知识内容。当用户消息匹配到知识条目的触发条件时，相关知识内容会自动注入到 SessionAgent 的上下文中，增强 AI 回复的准确性和专业性。系统通过调度作业 world_book_update 定期同步企业数据源，保持知识库的时效性。")

    # 3.18 消息去重与幂等性设计
    add_heading_styled(doc, "3.18  消息去重与幂等性设计", level=2)
    add_para(doc, "消息去重是 Emily 消息接入层的关键机制。薄插件（emily_agent）使用 SHA256 消息指纹对入站消息进行去重，防止同一条消息被重复处理。指纹由消息 ID、发送者 ID 和消息内容组合计算得出，确保去重的准确性。")
    add_para(doc, "在业务层，Emily 通过以下机制保证操作的幂等性：事件的创建使用消息 ID 作为幂等键（source_message_id），同一条消息不会重复创建事件；数据库操作使用事务确保原子性；WorkItem 的状态机限制了非法的状态转换（如从 DONE 不能再转为 EXECUTING），防止重复执行。这些机制共同保证了系统在消息重试、网络抖动等异常场景下的数据一致性。")

    # 3.19 出站事件总线设计
    add_heading_styled(doc, "3.19  出站事件总线设计", level=2)
    add_para(doc, "OutboundEventBus 是 Emily 的出站消息分发核心，基于 asyncio.Queue 实现异步发布/订阅模式。当 PipelineBUS 完成 WorkItem 执行后，生成的 ReplyMessage 通过 OutboundEventBus 发布，SSE 监听器（薄插件端）订阅并转发到 IM 平台。")
    add_para(doc, "NodeEventBus 用于全景节点图的实时变更通知，当节点状态发生变更（创建/更新/激活/废弃/进度更新）时，通过 NodeEventBus 发布事件，前端通过 SSE 端点（/api/v1/node/events）实时接收更新。两个事件总线独立运行，互不干扰，确保了消息处理的并行性和可靠性。")
    add_para(doc, "事件总线的设计使得消息生产者和消费者完全解耦。EmilyCore 作为生产者只负责发布事件，不需要关心消费者是否存在或如何处理。薄插件作为消费者独立订阅和处理，可以根据需要启动或停止消费而不影响生产端的正常运行。")

    # ==================== 4. 数据库设计 ====================
    doc.add_page_break()
    add_heading_styled(doc, "4  数据库设计", level=1)

    add_heading_styled(doc, "4.1  数据库总体架构", level=2)
    add_para(doc, "系统使用独立 PostgreSQL 数据库（emily 库），通过 SQLAlchemy 2.0 DeclarativeBase ORM 映射，共 53 张数据表。数据库运行在独立的 Docker 容器（emily-postgres）中，通过 Docker 内部网络与应用容器通信。")
    add_para(doc, "所有表使用 UUID 主键（通过 _new_uuid() 生成），业务编号采用 {PREFIX}-YYYYMMDD-{uuid8} 格式（如 EVT-20260710-a1b2c3d4）。时间戳统一使用 ISO8601 字符串（如 2026-07-10T09:30:00+08:00），时区为北京时间（+08:00）。连接池配置为 pool_size=5、max_overflow=10、pool_pre_ping=True、pool_recycle=3600。")
    add_para(doc, "数据库初始化通过 init_db() 函数完成，使用 Base.metadata.create_all() 幂等建表。对于已存在的表，不会自动加列或修改索引。emily-core 启动时通过 _ensure_columns() 自动检查 information_schema，补齐 _PENDING_COLUMNS 映射中注册的缺失列。新增列需在 _PENDING_COLUMNS 中预先注册。")

    add_heading_styled(doc, "4.2  核心数据表", level=2)
    add_para(doc, "以下列出系统中的核心数据表及其职责：")

    add_table(doc,
        ["表名", "功能说明", "列数"],
        [
            ["users", "人员信息（系统身份 + 人事档案，合并原 employees 表）", "20"],
            ["user_im_bindings", "IM 平台账号 → 系统用户绑定（QQ→User 映射）", "8"],
            ["conversations", "对话会话（群聊/私聊，Session 生命周期载体）", "10"],
            ["messages", "全量通讯记录（入站+出站+前导消息，22个字段）", "22"],
            ["message_attachments", "消息附件关联中间表（一消息多附件）", "10"],
            ["projects", "项目基本条件与生命周期阶段管理", "13"],
            ["events", "项目事件记录（如：放线完成、材料进场）", "16"],
            ["tasks", "任务管理（负责人/截止日期/状态追踪）", "14"],
            ["meetings", "会议记录（参与人/摘要/关联文件）", "20"],
            ["files", "文件存储（版本链+保密级别+附件来源追溯）", "26"],
            ["company_info", "参建公司基础信息（社会信用代码等）", "9"],
            ["project_indicator_details", "项目指标明细（值/单位/来源/约束性）", "12"],
            ["scheduler_jobs", "系统调度作业（ONCE/CRON/INTERVAL 三种类型）", "16"],
            ["scheduler_executions", "调度执行记录（PENDING/RUNNING/SUCCESS/FAILED）", "8"],
            ["permission_def", "权限码定义（资源路径编码）", "10"],
            ["permission_grants", "授权记录（AUTO/TEMP/PERMANENT + 有效期）", "15"],
            ["permission_requests", "权限申请审批（PENDING/APPROVED/REJECTED）", "19"],
            ["permission_audit_log", "授权审计日志（仅 INSERT，不可篡改）", "11"],
            ["public_field_registry", "公开字段白名单登记（模型-字段级）", "6"],
            ["agent_reasoning_logs", "Agent 推理记录（迭代/耗时/SOP匹配）", "14"],
            ["llm_interaction_logs", "LLM 调用日志（模型/token/延迟/类型）", "16"],
            ["tool_call_logs", "工具调用日志（参数/结果/耗时/成功标志）", "12"],
            ["hook_execution_logs", "Hook 执行日志（决策/耗时/挂载点）", "10"],
            ["sop_business_flows", "SOP 业务流注册（权限/适用范围/版本）", "20"],
            ["sop_permission_bindings", "SOP↔权限组 M:N 绑定", "6"],
            ["permission_groups", "权限组（按参与方类型/部门/层级）", "15"],
        ],
        col_widths=[4.0, 7.5, 1.5]
    )

    add_heading_styled(doc, "4.3  数据关系图", level=2)
    add_ascii_diagram(doc, """
  ┌──────────┐     ┌────────────────┐     ┌───────────┐
  │  users   │────▶│user_im_bindings│     │ projects  │
  └────┬─────┘     └────────────────┘     └──┬────────┘
       │                                      │
       │    ┌──────────────┐    ┌────────────┐│
       │    │conversations │───▶│ messages   ││
       │    └──────────────┘    └────────────┘│
       │                                      │
       │         ┌────────┐  ┌────────┐  ┌─────────┐
       └────────▶│ events │  │ tasks  │  │meetings │
                 └───┬────┘  └───┬────┘  └───┬─────┘
                     │           │           │
                     └───────────┼───────────┘
                                 │ project_id (FK)
                                 ▼
                          ┌──────────┐
                          │ projects │
                          └──────────┘

  ┌───────────────────┐     ┌─────────────────────┐
  │scheduler_jobs     │────▶│scheduler_executions  │
  └───────────────────┘     └─────────────────────┘

  ┌───────────────────┐     ┌─────────────────────┐
  │permission_def     │     │permission_grants     │
  └───────────────────┘     └─────────────────────┘

  ┌───────────────────┐     ┌─────────────────────┐
  │sop_business_flows │────▶│sop_permission_bindings│
  └───────────────────┘     └─────────────────────┘
""", "图 4-1  核心数据表关系图")

    # ==================== 5. 接口设计 ====================
    doc.add_page_break()
    add_heading_styled(doc, "5  接口设计", level=1)

    add_heading_styled(doc, "5.1  HTTP API 接口", level=2)
    add_para(doc, "emily-core 容器通过 FastAPI 暴露以下 HTTP API 接口，所有接口路径以 /api/v1 为前缀：")

    add_table(doc,
        ["方法", "路径", "功能说明", "认证"],
        [
            ["POST", "/api/v1/message/send", "薄插件转发入站消息", "API Key"],
            ["POST", "/api/v1/session/terminate", "强制终止指定 Session", "API Key"],
            ["GET", "/api/v1/events/outbound", "SSE 出站事件流（长连接）", "API Key"],
            ["GET", "/api/v1/health", "健康检查", "无"],
            ["GET", "/api/v1/permission/def", "查询权限码定义", "API Key"],
            ["POST", "/api/v1/permission/grant", "授权权限", "API Key + Admin"],
            ["GET", "/api/v1/skills/list", "列出已注册技能", "API Key"],
            ["GET/POST", "/api/v1/node/*", "全景节点图 CRUD 接口", "API Key"],
            ["GET", "/api/v1/node/events", "节点变更 SSE 事件流", "API Key"],
            ["GET/POST", "/api/v1/evolution/*", "进化管理接口", "API Key + Admin"],
            ["GET/POST", "/api/v1/meta-cognition/*", "元认知管理接口", "API Key + Admin"],
        ],
        col_widths=[1.5, 4.0, 4.0, 2.5]
    )

    add_heading_styled(doc, "5.2  SSE 事件推送接口", level=2)
    add_para(doc, "系统通过 Server-Sent Events（SSE）实现服务端向客户端的实时事件推送。出站事件流（/api/v1/events/outbound）将 AI 回复推送给薄插件，再由薄插件通过 AstrBot 发送到 QQ。节点事件流（/api/v1/node/events）推送全景节点图的实时变更通知，用于前端实时更新。")
    add_para(doc, "SSE 底层基于 asyncio.Queue 实现发布/订阅模式，由 OutboundEventBus 和 NodeEventBus 统一管理。OutboundEventBus 支持多消费者同时监听，每个消费者独立消费队列。事件格式为标准 SSE data 字段，包含事件类型、内容和时间戳。")

    add_heading_styled(doc, "5.3  标准协议对象", level=2)
    add_para(doc, "跨容器通信的标准协议对象是 IM 适配器与 Emily Core 之间的跨平台合约，Core 只依赖这些对象，不依赖原始 IM 事件。以下列出核心协议对象：")

    add_table(doc,
        ["对象名", "定义文件", "用途", "关键字段"],
        [
            ["StandardMessage", "adapters/standard/message.py", "入站统一消息", "message_id, platform, conversation_type,\nsender_id, content, is_at_bot, attachments"],
            ["ReplyMessage", "adapters/standard/reply.py", "出站统一回复", "conversation_id, content, reply_to_message_id,\nreferences, file_paths"],
            ["RouteDecision", "adapters/standard/route_decision.py", "适配器层接管判断", "takeover, mode, confidence,\nshould_reply, reason"],
            ["RouteResult", "adapters/standard/result.py", "结构化参数载体", "intent, project_id, confidence, data"],
            ["HandlerResult", "adapters/standard/result.py", "Application handler 输出", "success, object_type, object_id,\nreply, error_code"],
            ["EventCommand", "adapters/standard/command.py", "事件创建参数信封", "project_id, title, event_type,\ncategory, description"],
            ["TaskCommand", "adapters/standard/command.py", "任务创建参数信封", "project_id, title, assignee_text,\ndue_date, description"],
        ],
        col_widths=[2.5, 3.5, 3.0, 4.0]
    )

    add_heading_styled(doc, "5.4  管道阶段接口", level=2)
    add_para(doc, "Pipeline 层使用抽象基类（ABC）定义各阶段的可替换引擎合约，确保 Mock 和 Real 实现的接口一致性：")

    add_table(doc,
        ["接口", "定义文件", "核心方法", "说明"],
        [
            ["AuthEngine", "pipeline/interfaces/auth.py", "authorize(user_id, route_decision)", "权限鉴权，输出 ALLOW/DENY"],
            ["WorkAgent", "pipeline/interfaces/planning.py", "plan(route_decision, context)", "生成 ExecutionPlan"],
            ["WorkAgent", "pipeline/interfaces/execution.py", "execute(plan, context)", "执行 PlanStep，输出 StepResult"],
            ["Guardian", "pipeline/interfaces/guardian.py", "review_step() / review_reply()", "步骤审核与出站审核"],
            ["RiskGrader", "pipeline/interfaces/risk.py", "grade(route_decision, operation_type)", "风险评估分级 L1/L2/L3"],
        ],
        col_widths=[2.5, 3.5, 3.5, 3.5]
    )

    add_para(doc, "这些接口的设计遵循依赖倒置原则：PipelineBUS 依赖抽象接口而非具体实现，使得每个节点可以独立替换为 Mock 或 Real 实现，无需修改管道代码。Mock 实现位于 pipeline/mocks/ 目录，Real 实现由 WorkItemAgent 统一提供。")

    # ==================== 6. 部署与运维设计 ====================
    doc.add_page_break()
    add_heading_styled(doc, "6  部署与运维设计", level=1)
    add_para(doc, "系统采用 Docker Compose 编排部署，所有容器通过 docker-compose-napcat.yml 配置文件管理。部署步骤如下：")

    steps = [
        "环境准备：安装 Docker（20.10+）和 Docker Compose（v2），确保端口 6099/8080/18080/5432 未被占用。服务器最低配置：CPU 4核、内存 8GB、磁盘 50GB。",
        "配置文件准备：编辑 emily-data/config/core_config.json 配置非机密运行时参数（LLM API 地址、数据库连接串、RAG 服务地址等），编辑 hook_config.json 配置 Hook 挂载规则。",
        "SOP 部署：将业务流手册（.md 文件）放置于 emily-data/sops/ 目录，SkillRegistry 在系统启动时自动扫描并索引，无需重启即可生效新 SOP。",
        "Prompt 模板：将 Agent system prompt 模板放置于 emily-data/prompts/ 目录，支持 session/workitem/guardian_step/guardian_reply/project 等多种模板。",
        "数据库初始化：emily-core 容器启动时自动执行 bootstrap.init()，通过 SQLAlchemy create_all() 幂等建表，通过 _ensure_columns() 自动补齐缺失列。",
        "启动服务：执行 docker compose -f docker-compose-napcat.yml up -d 启动全部五个容器。",
        "健康检查：访问 GET /api/v1/health 确认 emily-core 服务就绪，检查数据库连接和 LLM/RAG 可用性。",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"({i}) {s}")

    add_para(doc, "运维注意事项：")
    notes = [
        "代码变更后需清除 __pycache__（Docker bind-mount 不触发 Python 重编译）并重启 emily-core 容器。",
        "数据库已有表不会自动 ALTER，新增列需在 _PENDING_COLUMNS 映射中注册，或手动执行 DDL。",
        "容器日志通过 docker logs --tail 100 emily-core 命令查看，支持实时跟踪。",
        "生产环境实战测试通过 emy-test CLI 工具进行，需使用 users 表中的真实用户 UUID，避免数据污染。",
        "全景节点批量操作通过 scripts/manage_nodes.py CLI 脚本执行，支持 create/update/activate/discard/progress/query 等命令。",
    ]
    for n in notes:
        add_para(doc, f"● {n}")

    # ==================== 7. 功能效果展示 ====================
    doc.add_page_break()
    add_heading_styled(doc, "7  功能效果展示", level=1)
    add_para(doc, "以下展示 Emily 系统的主要功能交互效果。图中为 QQ 即时通讯平台中的实际对话截图，展示了用户与 Emily 的自然语言交互过程和系统回复效果。")

    add_heading_styled(doc, "7.1  事件记录功能", level=2)
    add_para(doc, "用户在 QQ 群中发送自然语言消息（如「帮我记录一个事件：样板段放线完成」），Emily 自动识别事件意图，提取事件类型、标题、项目等关键信息，创建事件记录并回复确认信息。系统支持多种事件类型的自动识别和分类。")
    add_placeholder_image(doc, "图 7-1  事件记录功能交互截图", height_cm=7.0)

    add_heading_styled(doc, "7.2  任务管理功能", level=2)
    add_para(doc, "用户通过自然语言创建任务（如「给张工分配任务：明天之前完成钢筋验收」），Emily 解析任务内容、负责人、截止日期等信息并录入系统。任务创建后系统自动关联项目，支持后续的任务查询和状态更新。")
    add_placeholder_image(doc, "图 7-2  任务管理功能交互截图", height_cm=7.0)

    add_heading_styled(doc, "7.3  知识库查询功能", level=2)
    add_para(doc, "用户提出业务相关问题（如「安全帽佩戴标准是什么」），Emily 通过 RAG 检索企业知识库并返回准确答案，附带知识来源引用（references），实现知识溯源。当主检索引擎不可用时，自动回退到本地关键词匹配。")
    add_placeholder_image(doc, "图 7-3  知识库查询功能交互截图", height_cm=7.0)

    add_heading_styled(doc, "7.4  全景节点图管理功能", level=2)
    add_para(doc, "用户通过自然语言创建、查询和更新项目全景节点图。Emily 以结构化文本回复节点树信息，支持里程碑→阶段→工序→交付物的四级树状结构。节点进度支持文本和语音两种方式更新。")
    add_placeholder_image(doc, "图 7-4  全景节点图管理功能交互截图", height_cm=7.0)

    add_heading_styled(doc, "7.5  SOP 业务流引导功能", level=2)
    add_para(doc, "用户触发 SOP 业务流后（如「我要提交一个设计变更申请」），Emily 按照 SOP 手册定义的步骤逐步引导用户完成操作，每步提供明确的指引、需要提交的材料和验收标准。完成后系统自动记录执行过程。")
    add_placeholder_image(doc, "图 7-5  SOP 业务流引导功能交互截图", height_cm=7.0)

    add_heading_styled(doc, "7.6  系统健康检查功能", level=2)
    add_para(doc, "管理员可通过 API 查看系统运行状态，包括各子系统健康度、数据库连接状态、LLM 服务可用性和调度任务执行情况。健康检查端点返回 JSON 格式的系统状态报告。")
    add_placeholder_image(doc, "图 7-6  系统健康检查功能截图", height_cm=7.0)

    add_heading_styled(doc, "7.7  权限管理功能", level=2)
    add_para(doc, "系统根据用户权限等级自动进行访问控制。不同权限等级的用户只能访问对应等级的 SOP 和数据。低权限用户执行受限操作时收到明确的权限不足提示，可申请临时权限。")
    add_placeholder_image(doc, "图 7-7  权限管理功能交互截图", height_cm=7.0)

    add_heading_styled(doc, "7.8  晨报与洞察推送功能", level=2)
    add_para(doc, "系统通过调度引擎在每日指定时间自动生成晨报，汇总前一天的事件、任务进展和异常情况，推送到指定群聊。洞察分析自动发现数据中的模式和异常，辅助决策。")
    add_placeholder_image(doc, "图 7-8  晨报与洞察推送功能截图", height_cm=7.0)

    return doc


def main():
    print("生成设计说明书 docx...")
    doc = build_document()
    output_path = ROOT / OUTPUT_FILE
    doc.save(str(output_path))

    total_paras = len(doc.paragraphs)
    estimated_pages = total_paras // 25 + 1
    print(f"已保存: {output_path}")
    print(f"段落总数: {total_paras}, 估计页数: ~{estimated_pages}")


if __name__ == "__main__":
    main()
