"""生成软著程序鉴别材料 docx —— 源代码前后各30页，每页≥50行。

用法: python scripts/gen_source_docx.py
输出: ./Emily_程序鉴别材料.docx
"""

import re
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---- 配置 ----
SOFTWARE_NAME = "Emily 企业公共大脑系统"
VERSION = "V0.7.0"
LINES_PER_PAGE = 50
FRONT_PAGES = 30
BACK_PAGES = 30
OUTPUT_FILE = "Emily_程序鉴别材料.docx"

ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "emily-core" / "emily_core"


def collect_source_files() -> list[Path]:
    """按层级排序收集所有 .py 文件（排除 __pycache__、scripts）。"""
    files = []
    for p in sorted(SOURCE_DIR.rglob("*.py")):
        rel = p.relative_to(SOURCE_DIR)
        parts = rel.parts
        if "__pycache__" in parts:
            continue
        if "scripts" in parts:
            continue
        files.append(p)
    return files


def concat_sources(files: list[Path]) -> list[str]:
    """拼接所有源文件，每个文件前加文件名标记行。返回行列表。"""
    all_lines: list[str] = []
    for f in files:
        rel = f.relative_to(ROOT)
        all_lines.append(f"# ---- 文件: {rel} ----")
        text = f.read_text(encoding="utf-8", errors="replace")
        for line in text.splitlines():
            all_lines.append(line)
        all_lines.append("")  # 文件间空行
    return all_lines


def sanitize_line(line: str) -> str:
    """脱敏：去除 API 密钥、密码等敏感信息。"""
    # 替换明显密码赋值
    line = re.sub(r'(password|passwd|pwd|secret|token|api_key)\s*=\s*["\'][^"\']+["\']',
                  r'\1 = "***REDACTED***"', line, flags=re.IGNORECASE)
    # 替换属性引用（如 self.config.llm_api_key → self.config.llm_api_key # REDACTED）
    line = re.sub(r'\.llm_api_key\b', '.llm_api_key  # REDACTED_VALUE', line)
    line = re.sub(r'\.api_key\b', '.api_key  # REDACTED_VALUE', line)
    # 替换连接串中的密码
    line = re.sub(r'(postgresql://[^:]+:)[^@]+(@)',
                  r'\1***REDACTED***\2', line)
    # 替换包含密钥值的环境变量
    line = re.sub(r'(EMILY_API_KEY|API_KEY)\s*=\s*["\'][^"\']+["\']',
                  r'\1 = "***REDACTED***"', line, flags=re.IGNORECASE)
    return line


def add_source_page(doc: Document, lines: list[str], page_num: int):
    """向 docx 添加一页源代码（50行），使用等宽字体。"""
    # 页眉 —— 在第一次调用时设置，使用 PAGE 域代码自动显示页码
    if page_num == 1:
        section = doc.sections[0]
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页眉文本 + PAGE 域代码
        run1 = hp.add_run(f"{SOFTWARE_NAME} {VERSION}  源程序代码  第 ")
        run1.font.size = Pt(9)
        run1.font.name = "宋体"

        # 插入 PAGE 域代码
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2 = hp.add_run()
        run2._element.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run3 = hp.add_run()
        run3._element.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run4 = hp.add_run()
        run4._element.append(fldChar2)

        run5 = hp.add_run(" 页")
        run5.font.size = Pt(9)
        run5.font.name = "宋体"

    # 源代码内容 —— 每行一个段落，等宽小字体
    for line in lines:
        line = sanitize_line(line)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)  # 紧凑行距

        run = p.add_run(line if line else " ")
        run.font.size = Pt(7.5)
        run.font.name = "Courier New"
        # 中文字体回退
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')

    # 分页符
    if page_num < FRONT_PAGES + BACK_PAGES:
        doc.add_page_break()


def main():
    print("1. 收集源代码文件...")
    files = collect_source_files()
    print(f"   找到 {len(files)} 个 .py 文件")

    print("2. 拼接源代码...")
    all_lines = concat_sources(files)
    print(f"   总行数: {len(all_lines)}")

    total_pages_needed = FRONT_PAGES + BACK_PAGES  # 60 页
    total_lines_needed = total_pages_needed * LINES_PER_PAGE  # 3000 行

    if len(all_lines) <= total_lines_needed:
        # 不足60页，全部提交
        print(f"   总行数不足 {total_lines_needed}，全部提交")
        front_lines = all_lines
        back_lines = []
    else:
        front_lines = all_lines[: FRONT_PAGES * LINES_PER_PAGE]
        back_lines = all_lines[-(BACK_PAGES * LINES_PER_PAGE):]
        print(f"   前 {FRONT_PAGES} 页: 第1-{len(front_lines)} 行")
        print(f"   后 {BACK_PAGES} 页: 第{len(all_lines)-len(back_lines)+1}-{len(all_lines)} 行")

    print("3. 生成 docx...")
    doc = Document()

    # 页面设置 —— A4、窄边距
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 封面
    for _ in range(8):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SOFTWARE_NAME)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("程序鉴别材料")
    run.font.size = Pt(18)
    run.font.name = "黑体"

    doc.add_paragraph()

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run(f"版本号: {VERSION}")
    run.font.size = Pt(14)
    run.font.name = "宋体"

    doc.add_page_break()

    # 前连续30页
    page_num = 1
    for i in range(FRONT_PAGES):
        start = i * LINES_PER_PAGE
        end = start + LINES_PER_PAGE
        page_lines = front_lines[start:end]
        if not page_lines:
            break
        # 不足50行时用空行补齐
        while len(page_lines) < LINES_PER_PAGE:
            page_lines.append("")
        add_source_page(doc, page_lines, page_num)
        page_num += 1

    # 后连续30页
    if back_lines:
        for i in range(BACK_PAGES):
            start = i * LINES_PER_PAGE
            end = start + LINES_PER_PAGE
            page_lines = back_lines[start:end]
            if not page_lines:
                break
            while len(page_lines) < LINES_PER_PAGE:
                page_lines.append("")
            add_source_page(doc, page_lines, page_num)
            page_num += 1

    # 移除最后一个多余的分页符
    # (add_source_page 会在非最后一页加分页符，最后一页不加)

    output_path = ROOT / OUTPUT_FILE
    doc.save(str(output_path))
    print(f"4. 已保存: {output_path}")
    print(f"   总页数: 封面1页 + 代码{page_num-1}页")


if __name__ == "__main__":
    main()
